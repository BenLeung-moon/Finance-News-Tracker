"""Render the executive summary as a Word document from SummaryProfile sections.

Layout:
- Narrative / story sections on page 1 (compact).
- Citations on page 2 when a citations section exists.
"""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from finance_news_tracker.config import Settings
from finance_news_tracker.profiles.base import SummaryProfile, SummarySection

GREY = RGBColor(0x66, 0x66, 0x66)
HK_TZ = ZoneInfo("Asia/Hong_Kong")


def _to_hk(value: datetime) -> datetime:
    if value.tzinfo is None:
        value = value.replace(tzinfo=timezone.utc)
    return value.astimezone(HK_TZ)


def _parse_datetime(value: Any) -> datetime | None:
    if isinstance(value, datetime):
        return value
    if not isinstance(value, str) or not value:
        return None
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return None


def _format_generated_time(generated_at: datetime) -> str:
    return f"{_to_hk(generated_at).strftime('%Y-%m-%d %H:%M')} HK time"


def _relevance_label(score: Any) -> str:
    try:
        value = int(score)
    except (TypeError, ValueError):
        value = 0
    if value >= 80:
        return "High"
    if value >= 60:
        return "Medium"
    return "Low"


def _display_date(item: dict[str, Any]) -> str:
    pub = item.get("published_at")
    if pub:
        parsed = _parse_datetime(pub)
        if parsed:
            return _to_hk(parsed).strftime("%Y-%m-%d")
        return str(pub)[:10]
    collected = item.get("collected_at")
    if collected:
        parsed = _parse_datetime(collected)
        if parsed:
            return f"~{_to_hk(parsed).strftime('%Y-%m-%d')} (first seen)"
        return f"~{str(collected)[:10]} (first seen)"
    return "date unknown"


def _add_hyperlink(paragraph, url: str, text: str, size: int = 9) -> None:
    """Append a clickable external hyperlink run to ``paragraph``."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))  # half-points
    rpr.append(sz)
    run.append(rpr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _tight(paragraph, before: int = 0, after: int = 4) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _tight(p, before=6, after=2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)


def _items_for_section(
    section: SummarySection,
    items: list[dict[str, Any]],
    summary_profile: SummaryProfile,
    *,
    for_docx: bool = False,
) -> list[dict[str, Any]]:
    if section.kind == "stories":
        limit = section.max_items or 7
        return items[:limit]
    if section.kind == "grouped_stories":
        wanted = set(section.category_ids)
        for cat, section_id in summary_profile.category_section_map.items():
            if section_id == section.id:
                wanted.add(cat)
        matched = [
            item
            for item in items
            if str(item.get("analysis_category") or item.get("category") or "")
            in wanted
        ]
        return matched[: (section.max_items or 5)]
    if section.kind == "citations":
        if for_docx and section.max_items_docx is not None:
            limit = section.max_items_docx
        else:
            limit = section.max_items or 10
        return items[:limit]
    return []


def _story_takeaway(item: dict[str, Any]) -> str:
    return (
        item.get("impact")
        or item.get("why_it_matters")
        or item.get("summary")
        or ""
    )


def _add_story(
    doc: Document,
    item: dict[str, Any],
    index: int,
    settings: Settings,
    fields: list[str],
) -> None:
    profile = settings.active_profile
    labels = profile.report_labels
    p = doc.add_paragraph()
    _tight(p, after=4)
    head = p.add_run(f"{index}. {item['title']}")
    head.bold = True

    meta_bits: list[str] = [profile.source_label(item["source"])]
    if "signal" in fields:
        meta_bits.append(profile.format_signal(str(item.get("signal", ""))))
    if "relevance" in fields:
        meta_bits.append(
            f"Relevance {_relevance_label(item.get('relevance_score'))}"
        )
    if "entity" in fields and item.get("entity"):
        meta_bits.append(f"{labels.entity_label}: {item['entity']}")
    tag = p.add_run(f"  ({' · '.join(meta_bits)})")
    tag.font.size = Pt(8)
    tag.font.color.rgb = GREY

    if "takeaway" in fields:
        takeaway = _story_takeaway(item)
        if takeaway:
            p.add_run().add_break()
            line = p.add_run(takeaway)
            line.font.size = Pt(9)
    if "impact" in fields and item.get("impact"):
        p.add_run().add_break()
        impact = p.add_run(f"{labels.impact_label}: {item['impact']}")
        impact.font.size = Pt(9)
    if "suggested_action" in fields and item.get("suggested_action"):
        p.add_run().add_break()
        action = p.add_run(
            f"{labels.suggested_action_label}: {item['suggested_action']}"
        )
        action.font.size = Pt(9)

    if "date" in fields or "url" in fields:
        p.add_run().add_break()
        if "date" in fields:
            date_run = p.add_run(f"Date: {_display_date(item)}  |  ")
            date_run.font.size = Pt(8)
            date_run.font.color.rgb = GREY
        if "url" in fields:
            _add_hyperlink(p, item["url"], "Read more", size=8)


def write_word_summary(
    llm_summary: dict[str, Any],
    items: list[dict[str, Any]],
    generated_at: datetime,
    out_path: Path,
    *,
    settings: Settings,
    max_stories: int = 7,
    citation_items: list[dict[str, Any]] | None = None,
    scoring_provider: str = "",
    scoring_model: str = "",
    analysis_provider: str = "",
    analysis_model: str = "",
) -> Path:
    profile = settings.active_profile
    labels = profile.report_labels
    summary_profile = profile.resolve_summary_profile()
    citations = citation_items if citation_items is not None else items
    doc = Document()

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(46)
        section.left_margin = section.right_margin = Pt(54)

    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(9)

    title = doc.add_paragraph()
    _tight(title, after=2)
    tr = title.add_run(labels.report_title)
    tr.bold = True
    tr.font.size = Pt(16)

    meta = doc.add_paragraph()
    _tight(meta, after=2)
    llm_bits = []
    if scoring_provider and scoring_model:
        llm_bits.append(f"Scoring: {scoring_provider}/{scoring_model}")
    if analysis_provider and analysis_model:
        llm_bits.append(f"Analysis: {analysis_provider}/{analysis_model}")
    llm_line = ("  |  " + "  |  ".join(llm_bits)) if llm_bits else ""
    mr = meta.add_run(
        f"Generated {_format_generated_time(generated_at)}  |  "
        f"Sources: {labels.sources_line}{llm_line}"
    )
    mr.font.size = Pt(8)
    mr.font.color.rgb = GREY

    citation_sections = [s for s in summary_profile.sections if s.kind == "citations"]
    body_sections = [s for s in summary_profile.sections if s.kind != "citations"]

    for section in body_sections:
        if section.kind == "narrative":
            field_name = section.narrative_field or summary_profile.narrative_field
            narrative = (
                llm_summary.get(field_name)
                or llm_summary.get("market_read")
                or summary_profile.fallback_narrative
                or labels.fallback_market_read
            )
            _heading(doc, section.title)
            p = doc.add_paragraph(str(narrative))
            _tight(p, after=6)
            continue

        if section.kind in {"stories", "grouped_stories"}:
            # Honor legacy max_stories for plain stories sections
            section_items = _items_for_section(
                section, items, summary_profile, for_docx=True
            )
            if section.kind == "stories":
                section_items = section_items[:max_stories]
            _heading(doc, section.title)
            if not section_items:
                p = doc.add_paragraph("No items in this section.")
                _tight(p, after=4)
                continue
            fields = section.item_fields or summary_profile.story_fields
            for i, item in enumerate(section_items, 1):
                _add_story(doc, item, i, settings, fields)
            continue

        if section.kind == "watchlist":
            _heading(doc, section.title)
            watchlist = llm_summary.get("watchlist") or (
                summary_profile.default_watchlist or labels.default_watchlist
            )
            for w in watchlist:
                p = doc.add_paragraph(str(w), style="List Bullet")
                _tight(p, after=2)
                for r in p.runs:
                    r.font.size = Pt(9)

    if citation_sections:
        doc.add_page_break()
        for section in citation_sections:
            _heading(doc, section.title)
            pool = _items_for_section(
                section, citations, summary_profile, for_docx=True
            )
            for item in pool:
                p = doc.add_paragraph(style="List Bullet")
                _tight(p, after=3)
                _add_hyperlink(p, item["url"], item["title"], size=9)
                tail = p.add_run(
                    f"  — {profile.source_label(item['source'])}, "
                    f"{_display_date(item)}"
                )
                tail.font.size = Pt(8)
                tail.font.color.rgb = GREY

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
